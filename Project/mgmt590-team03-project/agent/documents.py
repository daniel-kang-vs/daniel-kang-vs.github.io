"""Multimodal document ingestion — extract text and images from uploaded files."""

from __future__ import annotations

import base64
import csv
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

TRIP_DATA_EXTENSIONS = {".parquet", ".csv"}
TEXT_EXTENSIONS = {".txt", ".md", ".json", ".yaml", ".yml", ".log"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc", ".rtf", *TEXT_EXTENSIONS}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
SUPPORTED_EXTENSIONS = TRIP_DATA_EXTENSIONS | DOCUMENT_EXTENSIONS | IMAGE_EXTENSIONS | {".xlsx", ".xls"}


@dataclass
class ParsedDocument:
    filename: str
    kind: str  # trip_data | text | image
    text: str = ""
    image_b64: Optional[str] = None
    mime_type: Optional[str] = None
    error: Optional[str] = None


@dataclass
class DocumentContext:
    """Combined context from one or more attachments for the LLM."""

    documents: list[ParsedDocument] = field(default_factory=list)
    combined_text: str = ""
    images: list[dict[str, str]] = field(default_factory=list)

    @property
    def filenames(self) -> list[str]:
        return [d.filename for d in self.documents]

    @property
    def has_trip_data(self) -> bool:
        return any(d.kind == "trip_data" for d in self.documents)

    @property
    def trip_data_files(self) -> list[ParsedDocument]:
        return [d for d in self.documents if d.kind == "trip_data"]

    def to_state(self) -> dict[str, Any]:
        return {
            "filenames": self.filenames,
            "combined_text": self.combined_text,
            "has_trip_data": self.has_trip_data,
            "image_count": len(self.images),
            "images": self.images,
        }


def _read_text_file(path: Path, *, max_chars: int = 50_000) -> str:
    text = path.read_text(encoding="utf-8", errors="replace")
    if len(text) > max_chars:
        return text[:max_chars] + f"\n\n[... truncated at {max_chars:,} characters ...]"
    return text


def _read_pdf(path: Path, *, max_chars: int = 50_000) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    text = "\n".join(parts).strip()
    if len(text) > max_chars:
        return text[:max_chars] + f"\n\n[... truncated at {max_chars:,} characters ...]"
    return text


def _read_docx(path: Path, *, max_chars: int = 50_000) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if len(text) > max_chars:
        return text[:max_chars] + f"\n\n[... truncated at {max_chars:,} characters ...]"
    return text


def _read_csv_preview(path: Path, *, max_rows: int = 30) -> str:
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = []
        for i, row in enumerate(reader):
            if i >= max_rows:
                rows.append(f"[... {max_rows} rows shown ...]")
                break
            rows.append(", ".join(row))
    return "\n".join(rows)


def _read_xlsx_preview(path: Path, *, max_rows: int = 30) -> str:
    import pandas as pd

    df = pd.read_excel(path, nrows=max_rows)
    return df.to_csv(index=False)


def _image_mime(suffix: str) -> str:
    mapping = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
    }
    return mapping.get(suffix.lower(), "image/png")


def parse_document(path: Path, original_name: str | None = None) -> ParsedDocument:
    """Parse a single uploaded file into text and/or image content."""
    name = original_name or path.name
    suffix = Path(name).suffix.lower()

    if suffix in TRIP_DATA_EXTENSIONS:
        return ParsedDocument(filename=name, kind="trip_data")

    if suffix in IMAGE_EXTENSIONS:
        raw = path.read_bytes()
        b64 = base64.standard_b64encode(raw).decode("ascii")
        return ParsedDocument(
            filename=name,
            kind="image",
            text=f"[Image attachment: {name}]",
            image_b64=b64,
            mime_type=_image_mime(suffix),
        )

    try:
        if suffix == ".pdf":
            text = _read_pdf(path)
        elif suffix in {".docx", ".doc"}:
            text = _read_docx(path)
        elif suffix in TEXT_EXTENSIONS:
            text = _read_text_file(path)
        elif suffix == ".csv":
            text = _read_csv_preview(path)
        elif suffix in {".xlsx", ".xls"}:
            text = _read_xlsx_preview(path)
        else:
            return ParsedDocument(
                filename=name,
                kind="text",
                error=f"Unsupported file type: {suffix or '(no extension)'}",
            )
        return ParsedDocument(filename=name, kind="text", text=text)
    except Exception as exc:
        return ParsedDocument(filename=name, kind="text", error=str(exc))


def parse_upload_bytes(data: bytes, filename: str) -> ParsedDocument:
    """Parse from in-memory upload (Streamlit file_uploader)."""
    suffix = Path(filename).suffix.lower()
    with io.BytesIO(data) as buf:
        if suffix in TRIP_DATA_EXTENSIONS:
            return ParsedDocument(filename=filename, kind="trip_data")

        tmp_suffix = suffix or ".bin"
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=tmp_suffix) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        try:
            return parse_document(tmp_path, original_name=filename)
        finally:
            tmp_path.unlink(missing_ok=True)


def build_document_context(docs: list[ParsedDocument]) -> DocumentContext:
    """Merge parsed documents into a single LLM context."""
    text_blocks: list[str] = []
    images: list[dict[str, str]] = []

    for doc in docs:
        if doc.error:
            text_blocks.append(f"### {doc.filename}\n[Parse error: {doc.error}]")
            continue
        if doc.kind == "trip_data":
            text_blocks.append(f"### {doc.filename}\n[Trip data file — applied as data overlay]")
            continue
        if doc.kind == "image" and doc.image_b64:
            images.append({"filename": doc.filename, "b64": doc.image_b64, "mime": doc.mime_type or "image/png"})
            if doc.text:
                text_blocks.append(f"### {doc.filename}\n{doc.text}")
            continue
        if doc.text.strip():
            text_blocks.append(f"### {doc.filename}\n{doc.text.strip()}")

    combined = "\n\n".join(text_blocks).strip()
    return DocumentContext(documents=docs, combined_text=combined, images=images)


def format_prompt_with_documents(user_text: str, ctx: DocumentContext | None) -> str:
    """Build the text portion of a user prompt including document content."""
    if not ctx or not ctx.combined_text:
        return user_text
    parts = []
    if user_text.strip():
        parts.append(f"User message:\n{user_text.strip()}")
    parts.append(f"Attached documents:\n{ctx.combined_text}")
    return "\n\n".join(parts)
